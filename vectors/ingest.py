"""
Ingestion script — indexes documents into Qdrant per agent.

Usage:
    # Ingest one agent's kb/ folder
    python ingest.py --agent ttu-online

    # Ingest all agents
    python ingest.py --all

    # Wipe collection and re-embed
    python ingest.py --agent ttu-online --full
    python ingest.py --all --full

    # Legacy: ingest a single file or URL directly (uses QDRANT_COLLECTION env)
    python ingest.py --file /path/to/doc.docx
    python ingest.py --url "https://texastechuniversity-my.sharepoint.com/..."
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from io import BytesIO
from pathlib import Path

import requests
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.table import Table as DocxTable
from docx.oxml.ns import qn as _qn
from transformers import AutoTokenizer
from qdrant_client.models import (
    Distance,
    PointStruct,
    SparseVector,
    SparseVectorParams,
    VectorParams,
)

from agents_config import AGENTS
from embeddings import embed_batch, sparse_embed_batch
from qdrant_store import get_client, reset_collection, upsert_points

CHUNK_TOKENS = 600
OVERLAP_TOKENS = 75
VECTOR_DIM = 384

_tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")


# ── Download ──────────────────────────────────────────────────────────────────

def download_sharepoint(url: str) -> tuple[bytes, str]:
    sep = "&" if "?" in url else "?"
    download_url = url + sep + "download=1"
    print("⬇️  Downloading from SharePoint...")
    resp = requests.get(download_url, allow_redirects=True, timeout=120)
    resp.raise_for_status()
    cd = resp.headers.get("Content-Disposition", "")
    filename = "document.docx"
    if "filename*=" in cd:
        filename = cd.split("filename*=")[-1].split("''")[-1].strip().strip('"\'')
    elif "filename=" in cd:
        filename = cd.split("filename=")[-1].strip().strip('"\'')
    print(f"   Got: {filename} ({len(resp.content):,} bytes)")
    return resp.content, filename


# ── Parse ─────────────────────────────────────────────────────────────────────

def extract_text_from_bytes(data: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        doc = DocxDocument(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if ext in (".txt", ".md"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {ext}")


def extract_text_from_file(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".docx":
        doc = DocxDocument(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if ext in (".txt", ".md"):
        return file_path.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported file type: {ext}")


# ── Rich docx extraction (images + step formatting) ──────────────────────────

STATIC_IMAGES_DIR = Path(__file__).parent / "static" / "images"
_IMAGE_PATTERN = re.compile(r'\[IMAGE:([^\]]+)\]')


def _split_chunk_images(chunk: str) -> tuple[str, list[str]]:
    """Strip [IMAGE:id] markers from chunk text, return (clean_text, image_ids)."""
    ids = _IMAGE_PATTERN.findall(chunk)
    clean = _IMAGE_PATTERN.sub('', chunk).strip()
    return clean, ids


def _iter_blocks(doc):
    """Yield paragraphs and tables from a Document in body order."""
    for child in doc.element.body:
        tag = child.tag
        if tag == _qn('w:p'):
            yield DocxParagraph(child, doc)
        elif tag == _qn('w:tbl'):
            yield DocxTable(child, doc)


def _has_numpr(para: DocxParagraph) -> bool:
    pPr = para._p.find(_qn('w:pPr'))
    return pPr is not None and pPr.find(_qn('w:numPr')) is not None


def extract_docx_rich(
    source: bytes | Path,
    source_prefix: str,
    images_dir: Path = STATIC_IMAGES_DIR,
) -> tuple[str, list[str]]:
    """
    Extract text from a .docx preserving step/bullet/table structure.
    Images are saved to images_dir as {source_prefix}_{n:03d}.{ext}.
    [IMAGE:id] markers are embedded in the returned text where images appear.
    Returns (text_with_markers, list_of_saved_image_ids).
    """
    if isinstance(source, bytes):
        doc = DocxDocument(BytesIO(source))
    else:
        doc = DocxDocument(str(source))

    images_dir.mkdir(parents=True, exist_ok=True)

    lines: list[str] = []
    img_counter = 0
    saved_ids: list[str] = []
    step_counter = 0

    for block in _iter_blocks(doc):
        if isinstance(block, DocxTable):
            step_counter = 0
            for row_idx, row in enumerate(block.rows):
                cells = [c.text.strip() for c in row.cells]
                row_text = " | ".join(c for c in cells if c)
                if not row_text:
                    continue
                if row_idx == 0:
                    lines.append(row_text)
                    lines.append("-" * min(len(row_text), 60))
                else:
                    lines.append(row_text)
            continue

        # ── Paragraph ────────────────────────────────────────────────────────
        para: DocxParagraph = block

        # Extract inline images from this paragraph
        for drawing in para._p.findall('.//' + _qn('w:drawing')):
            blip = drawing.find('.//' + _qn('a:blip'))
            if blip is None:
                continue
            rId = blip.get(_qn('r:embed'))
            if not rId or rId not in doc.part.related_parts:
                continue
            img_part = doc.part.related_parts[rId]
            ext = img_part.content_type.split('/')[-1].lower()
            if ext == 'jpeg':
                ext = 'jpg'
            if ext not in {'png', 'jpg', 'gif', 'bmp', 'tiff', 'webp'}:
                continue
            img_counter += 1
            img_id = f"{source_prefix}_{img_counter:03d}.{ext}"
            (images_dir / img_id).write_bytes(img_part.blob)
            saved_ids.append(img_id)
            lines.append(f"[IMAGE:{img_id}]")

        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if "List Number" in style_name or (
            "List Paragraph" in style_name and _has_numpr(para)
        ):
            step_counter += 1
            lines.append(f"Step {step_counter}: {text}")
        elif "List Bullet" in style_name or (
            "List Paragraph" in style_name and not _has_numpr(para)
        ):
            step_counter = 0
            lines.append(f"- {text}")
        elif "Heading" in style_name:
            step_counter = 0
            lines.append(f"\n{text}")
        else:
            step_counter = 0
            lines.append(text)

    return "\n".join(lines), saved_ids


# ── Chunk ─────────────────────────────────────────────────────────────────────

def _token_count(text: str) -> int:
    return len(_tokenizer.encode(text, add_special_tokens=False))


def _split_into_units(text: str) -> list[str]:
    units: list[str] = []
    for para in re.split(r"\n{2,}", text):
        para = para.strip()
        if not para:
            continue
        if _token_count(para) <= CHUNK_TOKENS:
            units.append(para)
        else:
            for line in para.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if _token_count(line) <= CHUNK_TOKENS:
                    units.append(line)
                else:
                    for sent in re.split(r"(?<=[.!?])\s+", line):
                        if sent.strip():
                            units.append(sent.strip())
    return units


def chunk_text(text: str) -> list[str]:
    units = _split_into_units(text)
    chunks: list[str] = []
    current_units: list[str] = []
    current_tokens = 0
    for unit in units:
        unit_tokens = _token_count(unit)
        if current_tokens + unit_tokens > CHUNK_TOKENS and current_units:
            chunks.append("\n\n".join(current_units))
            overlap: list[str] = []
            overlap_tokens = 0
            for u in reversed(current_units):
                t = _token_count(u)
                if overlap_tokens + t <= OVERLAP_TOKENS:
                    overlap.insert(0, u)
                    overlap_tokens += t
                else:
                    break
            current_units = overlap
            current_tokens = overlap_tokens
        current_units.append(unit)
        current_tokens += unit_tokens
    if current_units:
        chunks.append("\n\n".join(current_units))
    return [c for c in chunks if _token_count(c) > 20]


# ── Core pipeline ─────────────────────────────────────────────────────────────

def _embed_and_upsert(
    chunks: list[str],
    payloads: list[dict],
    collection_name: str,
    full: bool,
) -> None:
    client = get_client()

    if full:
        try:
            client.delete_collection(collection_name)
        except Exception:
            pass
        client.create_collection(
            collection_name=collection_name,
            vectors_config={"dense": VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)},
            sparse_vectors_config={"sparse": SparseVectorParams()},
        )
        print(f"  🗑️  Collection '{collection_name}' wiped and recreated.")
    else:
        existing = [c.name for c in client.get_collections().collections]
        if collection_name not in existing:
            client.create_collection(
                collection_name=collection_name,
                vectors_config={"dense": VectorParams(size=VECTOR_DIM, distance=Distance.COSINE)},
                sparse_vectors_config={"sparse": SparseVectorParams()},
            )

    print(f"  ⚙️  Dense + sparse embeddings for {len(chunks)} chunks...")
    dense_vectors = embed_batch(chunks)
    sparse_vectors = sparse_embed_batch(chunks)

    points = [
        PointStruct(
            id=i,
            vector={
                "dense": dense_vectors[i],
                "sparse": SparseVector(
                    indices=sparse_vectors[i][0],
                    values=sparse_vectors[i][1],
                ),
            },
            payload=payloads[i],
        )
        for i in range(len(chunks))
    ]
    client.upsert(collection_name=collection_name, points=points, wait=True)


def ingest_to_collection(
    text: str,
    source_name: str,
    collection_name: str,
    full: bool = False,
) -> int:
    """Embed text and store in a named collection.
    Automatically strips [IMAGE:id] markers from chunk text and stores ids in payload."""
    raw_chunks = chunk_text(text)
    clean_chunks: list[str] = []
    payloads: list[dict] = []
    for i, raw in enumerate(raw_chunks):
        clean, img_ids = _split_chunk_images(raw)
        clean_chunks.append(clean)
        payloads.append({
            "text": clean,
            "source": source_name,
            "chunk_index": i,
            "image_ids": img_ids,
        })
    _embed_and_upsert(clean_chunks, payloads, collection_name, full)
    return len(clean_chunks)


def run_ingestion_from_bytes(data: bytes, source_name: str, collection_name: str, full: bool = True) -> int:
    """Ingest from raw bytes (used by KB URL watcher). Uses rich extraction for .docx."""
    if source_name.lower().endswith(".docx"):
        slug = re.sub(r'[^a-z0-9]+', '_', Path(source_name).stem.lower()).strip('_')
        text, img_ids = extract_docx_rich(data, slug)
        print(f"✂️  Rich extraction: {len(img_ids)} images saved")
    else:
        text = extract_text_from_bytes(data, source_name)
    return ingest_to_collection(text, source_name, collection_name, full=full)


def run_ingestion(text: str, source_name: str) -> int:
    """Legacy plain-text ingestion kept for backward compat."""
    collection = os.getenv("QDRANT_COLLECTION", "ttu_kb")
    return ingest_to_collection(text, source_name, collection, full=True)


# ── Agent ingestion ───────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".docx", ".txt", ".md"}


def ingest_agent(agent_name: str, full: bool = False) -> None:
    if agent_name not in AGENTS:
        print(f"❌ Unknown agent '{agent_name}'. Available: {list(AGENTS.keys())}", file=sys.stderr)
        sys.exit(1)

    cfg = AGENTS[agent_name]
    collection = cfg["collection"]
    kb_folder = Path(cfg["kb_folder"])

    # Collect local files if the folder exists and has documents
    files = []
    if kb_folder.exists():
        files = [f for f in kb_folder.iterdir() if f.suffix.lower() in SUPPORTED_EXTENSIONS]

    # Fall back to kb_url if no local files
    if not files:
        kb_url = cfg.get("kb_url", "")
        if not kb_url:
            print(f"⚠️  No local files in {kb_folder} and no kb_url configured for '{agent_name}'.")
            return
        print(f"\n🤖 Agent: {cfg['display_name']} ({agent_name})")
        print(f"   Collection: {collection}")
        print(f"   No local files found — downloading from OneDrive URL...")
        data, filename = download_sharepoint(kb_url)
        count = run_ingestion_from_bytes(data, filename, collection, full=full)
        print(f"\n  ✅ Indexed {count} chunks into '{collection}'\n")
        return

    print(f"\n🤖 Agent: {cfg['display_name']} ({agent_name})")
    print(f"   Collection: {collection}")
    print(f"   Files: {len(files)}")

    all_chunks: list[str] = []
    all_payloads: list[dict] = []

    for f in sorted(files):
        print(f"\n  📄 {f.name}")
        if f.suffix.lower() == ".docx":
            slug = re.sub(r'[^a-z0-9]+', '_', f.stem.lower()).strip('_')
            text, img_ids = extract_docx_rich(f, slug)
            print(f"     🖼️  {len(img_ids)} images extracted → static/images/")
        else:
            text = extract_text_from_file(f)

        raw_chunks = chunk_text(text)
        print(f"     ✂️  {len(raw_chunks)} chunks")

        for i, raw in enumerate(raw_chunks):
            clean, chunk_imgs = _split_chunk_images(raw)
            all_chunks.append(clean)
            all_payloads.append({
                "text": clean,
                "source": f.name,
                "chunk_index": i,
                "image_ids": chunk_imgs,
            })

    _embed_and_upsert(all_chunks, all_payloads, collection, full)
    print(f"\n  ✅ Indexed {len(all_chunks)} total chunks into '{collection}'\n")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="Ingest documents into Qdrant.")
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--agent", metavar="NAME", help=f"Ingest one agent. Choices: {list(AGENTS.keys())}")
    mode.add_argument("--all", action="store_true", help="Ingest all agents.")
    mode.add_argument("--file", metavar="PATH", help="Legacy: ingest a single local file.")
    mode.add_argument("--url", metavar="URL", help="Legacy: ingest from a SharePoint URL.")
    parser.add_argument("--full", action="store_true", help="Wipe collection before ingesting.")
    args = parser.parse_args()

    if args.all:
        for name in AGENTS:
            ingest_agent(name, full=args.full)
        print("🎉 All agents ingested.")

    elif args.agent:
        ingest_agent(args.agent, full=args.full)

    elif args.url:
        data, filename = download_sharepoint(args.url)
        print(f"📄 Parsing: {filename}")
        text = extract_text_from_bytes(data, filename)
        count = run_ingestion(text, filename)
        print(f"🚀 Indexed {count} chunks into '{os.getenv('QDRANT_COLLECTION', 'ttu_kb')}'")

    else:  # --file
        file_path = Path(args.file).resolve()
        if not file_path.exists():
            print(f"File not found: {file_path}", file=sys.stderr)
            sys.exit(1)
        print(f"📄 Reading: {file_path}")
        text = extract_text_from_file(file_path)
        count = run_ingestion(text, file_path.name)
        print(f"🚀 Indexed {count} chunks into '{os.getenv('QDRANT_COLLECTION', 'ttu_kb')}'")


if __name__ == "__main__":
    main()
